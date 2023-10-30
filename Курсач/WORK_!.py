import docx


def replace_keywords_in_docx(docx_path, keyword_dict, output_path):
    doc = docx.Document(docx_path)
    slovar = keyword_dict

    for paragraph in doc.paragraphs:
        for k, v in slovar.items():
            if k in paragraph.text:
                paragraph.text = paragraph.text.replace(k, v)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for k, v in slovar.items():
                    if k in cell.text:
                        cell.text = cell.text.replace(k, v)
    doc.save(output_path)


# Пример использования
keyword_dict = {
    "${poln1}": "Значение1",
    "${naimenovanie1}": "Значение2",
    "${utv1}": "Значение3",
    "${opis1}": "Значение4",
    "${podpis1}": "Значение5",
    "${inic1}": "Значение6",
    "${mykey1}": "Ключ1",
    "${mykey2}": "Ключ2",
    "${mykey3}": "Ключ3"
}
# файл с ключами
input_file = "Приложение Д.docx"
# название итогового файта
output_file = "ITOG1.docx"

replace_keywords_in_docx(input_file, keyword_dict, output_file)
